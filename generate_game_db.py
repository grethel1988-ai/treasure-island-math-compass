import os
import sys
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re

def set_run_font(run, font_name="微軟正黑體", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    # Set font for East Asian text
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_formatted_text(p, text, font_name="微軟正黑體", size_pt=11, bold=False, italic=False, color_rgb=None):
    # Regex to match mixed fraction (e.g. 1又2/5) or simple fraction (e.g. 7/5)
    fraction_re = re.compile(r'(\d+又\d+/\d+|\d+/\d+)')
    parts = fraction_re.split(text)
    for part in parts:
        if not part:
            continue
        if '/' in part:
            if '又' in part:
                match = re.match(r'(\d+)又(\d+)/(\d+)', part)
                integer_part = match.group(1)
                num = match.group(2)
                den = match.group(3)
                xml = (
                    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                    f'<m:r><m:t>{integer_part}</m:t></m:r>'
                    '<m:f>'
                    f'<m:num><m:r><m:t>{num}</m:t></m:r></m:num>'
                    f'<m:den><m:r><m:t>{den}</m:t></m:r></m:den>'
                    '</m:f>'
                    '</m:oMath>'
                )
            else:
                match = re.match(r'(\d+)/(\d+)', part)
                num = match.group(1)
                den = match.group(2)
                xml = (
                    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                    '<m:f>'
                    f'<m:num><m:r><m:t>{num}</m:t></m:r></m:num>'
                    f'<m:den><m:r><m:t>{den}</m:t></m:r></m:den>'
                    '</m:f>'
                    '</m:oMath>'
                )
            p._p.append(parse_xml(xml))
        else:
            run = p.add_run(part)
            set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic, color_rgb=color_rgb)

def create_document():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("國小數學第三學習階段遊戲題庫")
    set_run_font(title_run, font_name="微軟正黑體", size_pt=18, bold=True, color_rgb=RGBColor(44, 62, 80))
    title_p.paragraph_format.space_after = Pt(6)

    # Subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("── 專案主題：金銀島：弗林特的寶藏（五至六年級能力指標 50 題） ──")
    set_run_font(sub_run, font_name="微軟正黑體", size_pt=12, italic=True, color_rgb=RGBColor(127, 140, 141))
    subtitle_p.paragraph_format.space_after = Pt(24)

    # Questions Data
    questions_data = [
        # n-III-1 (3 questions)
        {
            "indicator": "指標編碼：n-III-1 ── 理解數的十進位的位值結構，並能據以延伸認識更大與更小的數。",
            "num": 1,
            "text": "「弗林特」船長在「金銀島」上埋藏了巨大的財富。大副「比利·蓬斯」在手記中記載：『若要開鎖，需解開十進位的秘密：一兆是幾個一億？一億是幾個一萬？』請幫忙解開這個謎題。",
            "options": [
                "(A) 一兆是 1000個一億；一億是 1000個一萬",
                "(B) 一兆是 10000個一億；一億是 10000個一萬",
                "(C) 一兆是 100000個一億；一億是 10000個一萬",
                "(D) 一兆是 10000個一億；一億是 1000個一萬"
            ]
        },
        {
            "indicator": "指標編碼：n-III-1 ── 理解數的十進位的位值結構，並能據以延伸認識更大與更小的數。",
            "num": 2,
            "text": "少年「吉姆」在「比利·蓬斯」的水手衣物箱底發現了一張古老航海圖，角落用極小的紅墨水標記著一個小數「0.0075」。請問這個數讀作什麼？其中的「7」和「5」分別在什麼位數？",
            "options": [
                "(A) 零點零零七五；百分位、千分位",
                "(B) 零點零零七五；千分位、萬分位",
                "(C) 零點零七十五；十分位、百分位",
                "(D) 零點零零七五；十分位、萬分位"
            ]
        },
        {
            "indicator": "指標編碼：n-III-1 ── 理解數的十進位的位值結構，並能據以延伸認識更大與更小的數。",
            "num": 3,
            "text": "在「弗林特」船長留下的藏寶秘記中，寫著他掠奪的西班牙「杜布隆」金幣與「幾尼」金幣的總數：由 3個兆、5個億和8個十萬組成。這個金幣數寫成數字是多少？",
            "options": [
                "(A) 300050800000",
                "(B) 300050080000",
                "(C) 3000050800000",
                "(D) 3000050080000"
            ]
        },

        # n-III-2 (2 questions)
        {
            "indicator": "指標編碼：n-III-2 ── 在具體情境中，解決三步驟以上之常見應用問題。",
            "num": 4,
            "text": "「比利·蓬斯」住在「班波提督亭」旅館，他帶了 500英鎊，向旅店老闆買了 3瓶每瓶 25英鎊的朗姆酒，和 4大包每包 45英鎊的乾糧。老闆因為他是長住客，結帳時便宜了 10英鎊。請問「比利」最後還剩下幾英鎊？",
            "options": [
                "(A) 235英鎊",
                "(B) 245英鎊",
                "(C) 255英鎊",
                "(D) 265英鎊"
            ]
        },
        {
            "indicator": "指標編碼：n-III-2 ── 在具體情境中，解決三步驟以上之常見應用問題。",
            "num": 5,
            "text": "「班·甘恩」在「金銀島」上採集了 240個野生椰子。他剔除了壞掉的 15個，其餘每 15個裝成一箱，準備賣給來訪的「伊斯班紐拉號」船員，每箱賣 250枚西班牙金幣。如果全部賣完，「班·甘恩」一共可以收入多少枚金幣？",
            "options": [
                "(A) 3500枚金幣",
                "(B) 3750枚金幣",
                "(C) 4000枚金幣",
                "(D) 56250枚金幣"
            ]
        },

        # n-III-3 (3 questions)
        {
            "indicator": "指標編碼：n-III-3 ── 認識因數、倍數、質數、最大公因數、最小公倍數的意義、計算與應用。",
            "num": 6,
            "text": "「吉姆」發現藏寶圖上標註了兩個關鍵地標「望遠鏡山」與「骷髏山」的加密刻度分別是 36 和 48。若要推算出寶藏大洞的真正位置，必須求出這兩個數的最大公因數與最小公倍數。",
            "options": [
                "(A) 最大公因數 6，最小公倍數 144",
                "(B) 最大公因數 12，最小公倍數 72",
                "(C) 最大公因數 12，最小公倍數 144",
                "(D) 最大公因數 24，最小公倍數 288"
            ]
        },
        {
            "indicator": "指標編碼：n-III-3 ── 認識因數、倍數、質數、最大公因數、最小公倍數的意義、計算與應用。",
            "num": 7,
            "text": "「伊斯班紐拉號」的廚房裡儲備了 45個蘋果和 60個梨子以防敗血症。獨腳廚師「西爾弗」想把它們平分裝進木箱中，每箱的蘋果要一樣多，梨子也要一樣多。請問最多可以裝成幾箱？此時每箱有幾個蘋果？",
            "options": [
                "(A) 最多 5箱；每箱有 9個蘋果",
                "(B) 最多 15箱；每箱有 3個蘋果",
                "(C) 最多 15箱；每箱有 4個蘋果",
                "(D) 最多 9箱；每箱有 5個蘋果"
            ]
        },
        {
            "indicator": "指標編碼：n-III-3 ── 認識因數、倍數、質數、最大公因數、最小公倍數的意義、計算與應用。",
            "num": 8,
            "text": "「金銀島」的南北兩端各有一座哨角，甲哨角每 12分鐘吹響一次，乙哨角每 18分鐘吹響一次。如果早上 7點兩哨角同時吹響，下一次同時吹響是上午幾點幾分？",
            "options": [
                "(A) 上午 7點 30分",
                "(B) 上午 7點 36分",
                "(C) 上午 7點 54分",
                "(D) 上午 8點 12分"
            ]
        },

        # n-III-4 (2 questions)
        {
            "indicator": "指標編碼：n-III-4 ── 理解約分、擴分、通分的意義，並應用於異分母分數的加減。",
            "num": 9,
            "text": "「吉姆」大約清點了「伊斯班紐拉號」上的水手，發現 24 名船員中有 18 名是暗中倒戈「西爾弗」的叛變海盜，叛變人數佔全部船員的 18/24。請把這個分數約分成最簡分數，並說明做法。",
            "options": [
                "(A) 9/12；分子和分母同除以 2",
                "(B) 6/8；分子和分母同除以 3",
                "(C) 3/4；分子和分母同除以最大公因數 6",
                "(D) 3/4；分子和分母同乘以 6"
            ]
        },
        {
            "indicator": "指標編碼：n-III-4 ── 理解約分、擴分、通分的意義，並應用於異分母分數的加減。",
            "num": 10,
            "text": "「吉姆」在「金銀島」上探險，第一天往東走了 5/6 哩，第二天往東北走了 3/8 哩。請問「吉姆」一共走了幾哩？算算看：5/6 + 3/8 = ( )。",
            "options": [
                "(A) 8/14 哩",
                "(B) 1又5/24 哩",
                "(C) 1又7/24 哩",
                "(D) 29/14 哩"
            ]
        },

        # n-III-5 (2 questions)
        {
            "indicator": "指標編碼：n-III-5 ── 理解整數相除的分數表示的意義。",
            "num": 11,
            "text": "「班·甘恩」與「吉姆」等共 4 人在山洞中分食 3 塊烤羊肉，每人分到的烤羊肉要一樣多。請問每個人可以分到幾塊烤羊肉？用分數怎麼表示？",
            "options": [
                "(A) 1/3 塊烤肉",
                "(B) 3/4 塊烤肉",
                "(C) 4/3 塊烤肉",
                "(D) 1又1/3 塊烤肉"
            ]
        },
        {
            "indicator": "指標編碼：n-III-5 ── 理解整數相除的分數表示的意義。",
            "num": 12,
            "text": "「西爾弗」將 7加侖的朗姆酒平均分給 5個忠實的部下。算算看「7 ÷ 5」，並求其假分數與帶分數商值。",
            "options": [
                "(A) 5/7；1又2/5",
                "(B) 7/5；1又2/5",
                "(C) 7/5；1又5/2",
                "(D) 5/7；1又2/7"
            ]
        },

        # n-III-6 (2 questions)
        {
            "indicator": "指標編碼：n-III-6 ── 理解分數乘法和除法的意義、計算與應用。",
            "num": 13,
            "text": "「伊斯班紐拉號」上的淡水桶中僅存 3/4 加侖的水。「吉姆」口渴喝了這瓶淡水的 2/3，請問「吉姆」喝了幾加侖的淡水？",
            "options": [
                "(A) 5/12 加侖",
                "(B) 1/2 加侖",
                "(C) 9/8 加侖",
                "(D) 1又1/12 加侖"
            ]
        },
        {
            "indicator": "指標編碼：n-III-6 ── 理解分數乘法和除法的意義、計算與應用。",
            "num": 14,
            "text": "「伊斯班紐拉號」在順風航行時，每小時前進 1又2/3 浬。如果持續以這個速度航行了 2又1/5 小時，一共航行了幾浬？（算式為：1又2/3 × 2又1/5 = ( )）",
            "options": [
                "(A) 2又2/15 浬",
                "(B) 3又2/3 浬",
                "(C) 3又11/15 浬",
                "(D) 4又2/15 浬"
            ]
        },

        # n-III-7 (2 questions)
        {
            "indicator": "指標編碼：n-III-7 ── 理解小數乘法和除法的意義，能做直式計算與應用。",
            "num": 15,
            "text": "一包乾糧重 0.25 公斤，「吉姆」在逃跑時帶了 3.2 包乾糧。請問這些乾糧共重幾公斤？",
            "options": [
                "(A) 0.08 公斤",
                "(B) 0.8 公斤",
                "(C) 8 公斤",
                "(D) 80 公斤"
            ]
        },
        {
            "indicator": "指標編碼：n-III-7 ── 理解小數乘法和除法的意義，能做直式計算與應用。",
            "num": 16,
            "text": "「吉姆」試圖用密碼盤解開「比利·蓬斯」留下的日誌，指針在 4.8 刻度上，需除以係數 1.2。算算看「4.8 ÷ 1.2 = ( )」，並選出關於小數點移動過程的正確敘述。",
            "options": [
                "(A) 0.4；除數與被除數的小數點同時往左移一位",
                "(B) 4；除數與被除數的小數點同時往右移一位",
                "(C) 4；只有除數的小數點需要往右移一位",
                "(D) 40；除數小數點往右移兩位，被除數往右移一位"
            ]
        },

        # n-III-8 (2 questions)
        {
            "indicator": "指標編碼：n-III-8 ── 理解以四捨五入取概數，並進行合理估算。",
            "num": 17,
            "text": "「弗林特」船長埋藏的各國金幣中，法國路易金幣共有 248915 枚。用四捨五入法把 248915 取概數到萬位是多少？取概數到千位是多少？",
            "options": [
                "(A) 240000；248000",
                "(B) 250000；249000",
                "(C) 250000；248000",
                "(D) 249000；250000"
            ]
        },
        {
            "indicator": "指標編碼：n-III-8 ── 理解以四捨五入取概數，並進行合理估算。",
            "num": 18,
            "text": "航海士測量「伊斯班紐拉號」的羅盤偏差常數大約是 3.14159。用四捨五入法取到小數第二位（百分位）是多少？",
            "options": [
                "(A) 3.1",
                "(B) 3.14",
                "(C) 3.15",
                "(D) 3.142"
            ]
        },

        # n-III-9 (2 questions)
        {
            "indicator": "指標編碼：n-III-9 ── 理解比例關係的意義，並能據以觀察、表述、計算與解題，如比率、比例尺、速度、基準量等。",
            "num": 19,
            "text": "「金銀島」附近的港口小鎮共有 600 個居民，其中婦女佔了 45%。請問婦女有多少人？這題的比率是多少？",
            "options": [
                "(A) 270人；比率是45%",
                "(B) 270人；比率是55%",
                "(C) 330人；比率是45%",
                "(D) 330人；比率是55%"
            ]
        },
        {
            "indicator": "指標編碼：n-III-9 ── 理解比例關係的意義，並能據以觀察、表述、計算與解題，如比率、比例尺、速度、基準量等。",
            "num": 20,
            "text": "在「弗林特」船長的藏寶圖上，比例尺是 1 : 50000。「吉姆」量出地圖上「比利·蓬斯」標記的紅「X」與「望遠鏡山」山頂的距離是 4公分，實際上的距離是多少公里？",
            "options": [
                "(A) 0.2公里",
                "(B) 2公里",
                "(C) 20公里",
                "(D) 200公里"
            ]
        },

        # n-III-10 (2 questions)
        {
            "indicator": "指標編碼：n-III-10 ── 嘗試將較複雜的情境或模式中的數量關係以算式正確表述，並據以推理或解題。",
            "num": 21,
            "text": "「班·甘恩」在島上馴養了野雞和野兔共 10 隻並關在籠子裡。數一數牠們的腳一共有 32 隻。請問籠子裡有幾隻野雞？幾隻野兔？",
            "options": [
                "(A) 野雞有6隻，野兔有4隻",
                "(B) 野雞有5隻，野兔有5隻",
                "(C) 野雞有4隻，野兔有6隻",
                "(D) 野雞有3隻，野兔有7隻"
            ]
        },
        {
            "indicator": "指標編碼：n-III-10 ── 嘗試將較複雜的情境或模式中的數量關係以算式正確表述，並據以推理或解題。",
            "num": 22,
            "text": "鄉紳「特里勞尼」今年 42 歲，「吉姆」今年 12 歲。幾年後，鄉紳的年齡會剛好是「吉姆」的 3 倍？",
            "options": [
                "(A) 2年後",
                "(B) 3年後",
                "(C) 4年後",
                "(D) 5年後"
            ]
        },

        # n-III-11 (2 questions)
        {
            "indicator": "指標編碼：n-III-11 ── 認識量的常用單位及其換算，並處理相關的應用問題。",
            "num": 23,
            "text": "「弗林特」船長在遺囑中提到，要把「金銀島」上 3公頃 的森林地分給大副「比利」。請問 3公頃 = （）公畝 = （）平方公尺。",
            "options": [
                "(A) 30；3000",
                "(B) 300；3000",
                "(C) 300；30000",
                "(D) 3000；300000"
            ]
        },
        {
            "indicator": "指標編碼：n-III-11 ── 認識量的常用單位及其換算，並處理相關的應用問題。",
            "num": 24,
            "text": "「班·甘恩」用石頭圍了一塊長方形的羊圈，長是 500公尺，寬是 400公尺。這塊羊圈的面積是多少公頃？也就是幾平方公里？",
            "options": [
                "(A) 2公頃；0.02平方公里",
                "(B) 20公頃；0.2平方公里",
                "(C) 20公頃；2平方公里",
                "(D) 200公頃；2平方公里"
            ]
        },

        # n-III-12 (2 questions)
        {
            "indicator": "指標編碼：n-III-12 ── 理解容量、容積和體積之間的關係，並做應用。",
            "num": 25,
            "text": "「伊斯班紐拉號」上的淡水儲藏鐵箱，其體積與容量需要換算：一立方公尺是多少立方公分？一公升是多少立方公分？",
            "options": [
                "(A) 100000；1000",
                "(B) 1000000；100",
                "(C) 1000000；1000",
                "(D) 10000；1000"
            ]
        },
        {
            "indicator": "指標編碼：n-III-12 ── 理解容量、容積和體積之間的關係，並做應用。",
            "num": 26,
            "text": "「吉姆」在「比利·蓬斯」的水手衣物箱裡找到一個無蓋木盒，外面的長寬高分別是 22公分、12公分和 11公分，壁厚是 1公分。這個木盒的容積是多少立方公分？可以裝多少毫升的淡水？",
            "options": [
                "(A) 1800立方公分；1800毫升",
                "(B) 2000立方公分；2000毫升",
                "(C) 2200立方公分；2200毫升",
                "(D) 2904立方公分；2904毫升"
            ]
        },

        # s-III-1 (2 questions)
        {
            "indicator": "指標編碼：s-III-1 ── 認識角柱、角錐、圓柱與圓錐的特徵，並了解其展開圖。",
            "num": 27,
            "text": "「班·甘恩」在東北方山洞裡，將黃金金條碼成了一個五角柱形狀。請問這個五角柱一共有幾個面？幾條邊？幾個頂點？",
            "options": [
                "(A) 5個面、10條邊、5個頂點",
                "(B) 7個面、15條邊、10個頂點",
                "(C) 7個面、10條邊、15個頂點",
                "(D) 6個面、15條邊、10個頂點"
            ]
        },
        {
            "indicator": "指標編碼：s-III-1 ── 認識角柱、角錐、圓柱與圓錐的特徵，並了解其展開圖。",
            "num": 28,
            "text": "「金銀島」上的老海盜祭壇殘留著一個六角錐石柱。請問這個六角錐一共有幾個面？幾條邊？幾個頂點？",
            "options": [
                "(A) 6個面、12條邊、6個頂點",
                "(B) 7個面、12條邊、7個頂點",
                "(C) 8個面、12條邊、8個頂點",
                "(D) 7個面、18條邊、12個頂點"
            ]
        },

        # s-III-2 (2 questions)
        {
            "indicator": "指標編碼：s-III-2 ── 理解扇形、圓心角與其面積的關係。",
            "num": 29,
            "text": "水手們在「伊斯班紐拉號」上慶祝平安抵達，切分一個圓形麵包，切出一個 120度圓心角的扇形給「吉姆」。這個扇形麵包是幾分之幾圓？",
            "options": [
                "(A) 二分之一（1/2）圓",
                "(B) 三分之一（1/3）圓",
                "(C) 四分之一（1/4）圓",
                "(D) 三分之二（2/3）圓"
            ]
        },
        {
            "indicator": "指標編碼：s-III-2 ── 理解扇形、圓心角與其面積的關係。",
            "num": 30,
            "text": "「弗林特」船長的密碼羅盤半徑是 12公分，其中一個指向寶藏入口的 1/4圓扇形區域。請問這個扇形區域的圓心角是多少度？面積是多少平方公分？（圓周率以 3.14 計算）",
            "options": [
                "(A) 45度；113.04平方公分",
                "(B) 90度；113.04平方公分",
                "(C) 90度；226.08平方公分",
                "(D) 90度；452.16平方公分"
            ]
        },

        # s-III-3 (2 questions)
        {
            "indicator": "指標編碼：s-III-3 ── 理解平行四邊形、三角形與梯形的面積公式與應用。",
            "num": 31,
            "text": "「吉姆」和夥伴們在「金銀島」上搭建了一間底 15公尺、高 8公尺的平行四邊形臨時木屋。這間木屋的佔地面積是多少平方公尺？",
            "options": [
                "(A) 23平方公尺",
                "(B) 46平方公尺",
                "(C) 60平方公尺",
                "(D) 120平方公尺"
            ]
        },
        {
            "indicator": "指標編碼：s-III-3 ── 理解平行四邊形、三角形與梯形的面積公式與應用。",
            "num": 32,
            "text": "「伊斯班紐拉號」上掛著一面三角形的英國國旗，底是 12公寸，面積是 48平方公寸。這面旗子的高是多少公寸？",
            "options": [
                "(A) 4公寸",
                "(B) 6公寸",
                "(C) 8公寸",
                "(D) 16公寸"
            ]
        },

        # s-III-4 (2 questions)
        {
            "indicator": "指標編碼：s-III-4 ── 理解角柱（含正方體、長方體）與圓柱的體積與表面積的計算方式。",
            "num": 33,
            "text": "「比利·蓬斯」收藏祕密地圖的長方體鐵盒，長 10公分、寬 6公分、高 5公分。它的表面積和體積各是多少？",
            "options": [
                "(A) 表面積140平方公分，體積300立方公分",
                "(B) 表面積280平方公分，體積150立方公分",
                "(C) 表面積280平方公分，體積300立方公分",
                "(D) 表面積300平方公分，體積280立方公分"
            ]
        },
        {
            "indicator": "指標編碼：s-III-4 ── 理解角柱（含正方體、長方體）與圓柱的體積與表面積的計算方式。",
            "num": 34,
            "text": "「吉姆」在「伊斯班紐拉號」的貨艙找到一個正方體彈藥箱，邊長是 8公寸。它的表面積和體積各是多少？",
            "options": [
                "(A) 表面積64平方公寸，體積512立方公寸",
                "(B) 表面積384平方公寸，體積384立方公寸",
                "(C) 表面積384平方公寸，體積512立方公寸",
                "(D) 表面積512平方公寸，體積384立方公寸"
            ]
        },

        # s-III-5 (2 questions)
        {
            "indicator": "指標編碼：s-III-5 ── 以簡單推理，理解幾何形體的性質。",
            "num": 35,
            "text": "「吉姆」用來繪製「金銀島」航線的三角板是一個等腰三角形。已知它的頂角是 50度，那它的兩個底角各是多少度？",
            "options": [
                "(A) 50度",
                "(B) 65度",
                "(C) 80度",
                "(D) 130度"
            ]
        },
        {
            "indicator": "指標編碼：s-III-5 ── 以簡單推理，理解幾何形體的性質。",
            "num": 36,
            "text": "「伊斯班紐拉號」上的四邊形船尾甲板，其內角和是多少度？我們可以用什麼簡單的方法，把四邊形切開來證明？",
            "options": [
                "(A) 180度；連接對角線分成 2 個三角形證明",
                "(B) 360度；連接對角線分成 2 個三角形證明",
                "(C) 360度；連接對角線分成 4 個三角形證明",
                "(D) 540度；連接對角線分成 3 個三角形證明"
            ]
        },

        # s-III-6 (2 questions)
        {
            "indicator": "指標編碼：s-III-6 ── 認識線對稱的意義與其推論。",
            "num": 37,
            "text": "「西爾弗」在研究海盜骷髏旗時，「吉姆」提到骷髏圖案是個『線對稱圖形』。請問什麼是『線對稱圖形』？請選出一個生活中的線對稱圖形例子。",
            "options": [
                "(A) 沿著一條線對折後，兩邊不會重合的圖形；如平行四邊形",
                "(B) 沿著一條線對折後，兩邊能完全重合的圖形；如蝴蝶",
                "(C) 沿著一個點旋轉後能完全重合的圖形；如風車",
                "(D) 面積一樣大但形狀不同的圖形；如拼圖"
            ]
        },
        {
            "indicator": "指標編碼：s-III-6 ── 認識線對稱的意義與其推論。",
            "num": 38,
            "text": "「吉姆」拿著三張不同形狀的地圖殘片：等腰三角形、正方形與圓形。請問它們一共有幾條對稱軸？",
            "options": [
                "(A) 1條；2條；4條",
                "(B) 1條；4條；無限多條",
                "(C) 3條；4條；8條",
                "(D) 2條；4條；無限多條"
            ]
        },

        # s-III-7 (2 questions)
        {
            "indicator": "指標編碼：s-III-7 ── 認識平面圖形縮放的意義與應用。",
            "num": 39,
            "text": "「吉姆」用放大鏡觀察藏寶圖上一個三角形的「骷髏山」標記，各邊長變為原來的 3倍。這是一個什麼圖？它的每個內角度數會變為原來的幾倍？",
            "options": [
                "(A) 3倍放大圖；3倍",
                "(B) 3倍放大圖；1倍（不變）",
                "(C) 9倍放大圖；3倍",
                "(D) 1/3倍縮小圖；1倍（不變）"
            ]
        },
        {
            "indicator": "指標編碼：s-III-7 ── 認識平面圖形縮放的意義與應用。",
            "num": 40,
            "text": "船長把長 10公分、寬 6公分的長方形「伊斯班紐拉號」艙位圖，縮小為 1/2 倍繪製在筆記本上。這張縮小圖的長、寬 and 面積各是多少？",
            "options": [
                "(A) 長 5公分、寬 3公分，面積 30平方公分",
                "(B) 長 5公分、寬 3公分，面積 15平方公分",
                "(C) 長 20公分、寬 12公分，面積 240平方公分",
                "(D) 長 5公分、寬 3公分，面積 7.5平方公分"
            ]
        },

        # r-III-1 (2 questions)
        {
            "indicator": "指標編碼：r-III-1 ── 理解各種計算規則（含分配律），並協助四則混合計算與應用解題。",
            "num": 41,
            "text": "「弗林特」船長給水手們發放獎勵，每人分得 35 枚西班牙金幣，一共有 99 人。算算看，怎麼算最快？ 99 × 35 = ( )。你運用了什麼計算規則？",
            "options": [
                "(A) 3465；乘法對加法的分配律",
                "(B) 3465；乘法對減法的分配律",
                "(C) 3500；乘法結合律",
                "(D) 3435；乘法交換律"
            ]
        },
        {
            "indicator": "指標編碼：r-III-1 ── 理解各種計算規則（含分配律），並協助四則混合計算與應用解題。",
            "num": 42,
            "text": "密碼羅盤的解密指針公式分別是 3.8 × 4.2 和 6.2 × 4.2，相加即為解密鎖的密碼。算算看，怎麼算最快？ 3.8 × 4.2 + 6.2 × 4.2 = ( )。",
            "options": [
                "(A) 40",
                "(B) 42",
                "(C) 42.4",
                "(D) 100"
            ]
        },

        # r-III-2 (2 questions)
        {
            "indicator": "指標編碼：r-III-2 ── 熟練數（含分數、小數）的四則混合計算。",
            "num": 43,
            "text": "醫生「李甫西」在調配藥劑時，先加入了 3/4 加侖的淨水，又加入了 1.25 × 2 加侖的藥液。算算看： 3/4 + 1.25 × 2 = ( )。分數與小數混合時怎麼計算？",
            "options": [
                "(A) 3.25（或 3又1/4）",
                "(B) 4（或 4）",
                "(C) 2.25（或 2又1/4）",
                "(D) 3.5（或 3又1/2）"
            ]
        },
        {
            "indicator": "指標編碼：r-III-2 ── 熟練數（含分數、小數）的四則混合計算。",
            "num": 44,
            "text": "「吉姆」試圖用公式解鎖「班·甘恩」在山洞中擺放的羅盤鎖： (2.4 - 1.2) ÷ 3/5 = ( )。請算算看結果是多少？",
            "options": [
                "(A) 0.72",
                "(B) 2",
                "(C) 2.4",
                "(D) 3"
            ]
        },

        # r-III-3 (2 questions)
        {
            "indicator": "指標編碼：r-III-3 ── 觀察情境或模式中的數量關係，並用文字或符號正確表述，協助推理與解題。",
            "num": 45,
            "text": "醫生「李甫西」的年齡比大副「比利·蓬斯」大 4 歲。如果用 x 表示醫生的年齡，「比利」的年齡要怎麼表示？",
            "options": [
                "(A) x + 4 歲",
                "(B) x - 4 歲",
                "(C) 4x 歲",
                "(D) x ÷ 4 歲"
            ]
        },
        {
            "indicator": "指標編碼：r-III-3 ── 觀察情境或模式中的數量關係，並用文字或符號正確表述，協助推理與解題。",
            "num": 46,
            "text": "港口的一把佩劍賣 x 枚銀幣，「吉姆」買了 3 把，付了 500 枚銀幣。商家要找回他多少枚銀幣？用含有 x 的算式表述。",
            "options": [
                "(A) 500 - 3x 元",
                "(B) 3x - 500 元",
                "(C) 500 - x 元",
                "(D) 500 ÷ 3x 元"
            ]
        },

        # d-III-1 (2 questions)
        {
            "indicator": "指標編碼：d-III-1 ── 報讀圓形圖，製作折線圖與圓形圖，並據以做簡單推論。",
            "num": 47,
            "text": "「伊斯班紐拉號」上的圖書櫃有各種航海書籍，圓形百分圖中顯示：航海誌佔 40%、天文與氣象書佔 30%、地理圖冊佔 20%、其他佔 10%。如果一共有 200本書，航海誌有幾本？",
            "options": [
                "(A) 40本",
                "(B) 80本",
                "(C) 100本",
                "(D) 120本"
            ]
        },
        {
            "indicator": "指標編碼：d-III-1 ── 報讀圓形圖，製作折線圖與圓形圖，並據以做簡單推論。",
            "num": 48,
            "text": "「吉姆」想在「伊斯班紐拉號」的圓形舵盤上畫出一個代表 25% 的『安全轉向舵角區』。在圓形圖上，這個區域的圓心角應該是多少度？",
            "options": [
                "(A) 45度",
                "(B) 90度",
                "(C) 120度",
                "(D) 180度"
            ]
        },

        # d-III-2 (2 questions)
        {
            "indicator": "指標編碼：d-III-2 ── 能從資料或圖表的資料數據，解決關於「可能性」的簡單問題。",
            "num": 49,
            "text": "「班·甘恩」把 8 顆黑石子 and 2 顆白石子放進袋子裡。「吉姆」不看袋子摸出一顆石子，摸到黑石子的可能性大還是白石子大？為什麼？",
            "options": [
                "(A) 白石子大；因為白石子比較顯眼",
                "(B) 黑石子大；因為黑石子數量比較多",
                "(C) 一樣大；因為都是隨機摸取",
                "(D) 無法比較；要看摸球的人是誰"
            ]
        },
        {
            "indicator": "指標編碼：d-III-2 ── 能從資料或圖表的資料數據，解決關於「可能性」的簡單問題。",
            "num": 50,
            "text": "「西爾弗」和船員們在甲板上玩擲骰子遊戲，丟一個公正的骰子（點數1~6），丟出奇數點和偶數點的可能性一樣大嗎？為什麼？",
            "options": [
                "(A) 不一樣大；因為奇數比較多",
                "(B) 一樣大；因為奇數和偶數的點數各有3個",
                "(C) 不一樣大；因為 6 是偶數中最大的",
                "(D) 一樣大；因為骰子只有 2 個面"
            ]
        }
    ]

    current_indicator = ""
    indicator_count = 0
    
    for q in questions_data:
        # Check if we need to write indicator heading
        if q["indicator"] != current_indicator:
            current_indicator = q["indicator"]
            indicator_count += 1
            
            # Add spacer
            if indicator_count > 1:
                spacer_p = doc.add_paragraph()
                spacer_p.paragraph_format.space_before = Pt(12)
                spacer_p.paragraph_format.space_after = Pt(12)
                # We can add a simple horizontal line
                spacer_run = spacer_p.add_run("─" * 40)
                set_run_font(spacer_run, font_name="微軟正黑體", size_pt=10, color_rgb=RGBColor(200, 200, 200))
                spacer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            ind_p = doc.add_paragraph()
            ind_run = ind_p.add_run(current_indicator)
            set_run_font(ind_run, font_name="微軟正黑體", size_pt=12, bold=True, color_rgb=RGBColor(41, 128, 185))
            ind_p.paragraph_format.space_before = Pt(12)
            ind_p.paragraph_format.space_after = Pt(6)
            
        # Add question
        q_p = doc.add_paragraph()
        num_run = q_p.add_run(f"{q['num']}. ")
        set_run_font(num_run, font_name="微軟正黑體", size_pt=11, bold=False, color_rgb=RGBColor(51, 51, 51))
        add_formatted_text(q_p, q['text'], font_name="微軟正黑體", size_pt=11, bold=False, color_rgb=RGBColor(51, 51, 51))
        q_p.paragraph_format.space_before = Pt(6)
        q_p.paragraph_format.space_after = Pt(4)
        
        # Add options
        for opt in q["options"]:
            opt_p = doc.add_paragraph()
            opt_p.paragraph_format.left_indent = Inches(0.4)
            opt_p.paragraph_format.space_before = Pt(0)
            opt_p.paragraph_format.space_after = Pt(2)
            # Prefix like "(A) "
            prefix = opt[:4]
            opt_text = opt[4:]
            prefix_run = opt_p.add_run(prefix)
            set_run_font(prefix_run, font_name="微軟正黑體", size_pt=10.5, color_rgb=RGBColor(70, 70, 70))
            add_formatted_text(opt_p, opt_text, font_name="微軟正黑體", size_pt=10.5, color_rgb=RGBColor(70, 70, 70))
            
    # Save document
    output_path = r"D:\國小數學能力指標-高年級-金銀島的密碼羅盤\第三學習階段遊戲題庫.docx"
    try:
        doc.save(output_path)
        print(f"Successfully generated document at: {output_path}")
    except PermissionError:
        print(f"ERROR: Cannot save file to {output_path}.")
        print("Please check if '第三學習階段遊戲題庫.docx' is open in Microsoft Word or another program, and close it before running the script again.")
        sys.exit(1)

if __name__ == "__main__":
    create_document()
